import csv
from docx import Document

def normalize(name):
    return name.replace(" ", "").replace("-", "").lower()

def export_table_csv_docx(rows, headers, csv_path, docx_path, title):
    with open(csv_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(headers)
        writer.writerows(rows)

    doc = Document()
    doc.add_heading(title, 0)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light List Accent 1"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    doc.save(docx_path)

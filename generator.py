import os
import json
import glob
import shutil
import subprocess
from datetime import datetime
import csv
from docx import Document

def normalize(name):
    return name.replace(" ", "").replace("-", "").lower()

def export_table_csv_docx(rows, headers, csv_path, docx_path, title):
    # Write CSV
    with open(csv_path, "w", newline="", encoding="utf-8") as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(headers)
        for row in rows:
            writer.writerow(row)
    print(f"✅ CSV table saved to: {csv_path}")

    # Write DOCX
    doc = Document()
    doc.add_heading(title, 0)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light List Accent 1'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    # Data rows
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    doc.save(docx_path)
    print(f"✅ Word DOCX table saved to: {docx_path}")

def generate_diagram(input_dir, output_base_dir, image_format="png", scale="2"):
    """
    Core generation logic. 
    Returns path to timestamped output subfolder containing .mmd, image, and tables.
    """

    print("\n🟢 AWS Visualization Generator Running")
    print(f"✅ Input folder: {input_dir}")
    print(f"✅ Base output folder: {output_base_dir}")
    print(f"✅ Image format: {image_format}")
    print(f"✅ Scale factor: {scale}\n")

    # ✅ Create timestamped output subfolder
    timestamp = datetime.now().strftime("%Y-%m-%d-%H%M%S")
    final_output_dir = os.path.join(output_base_dir, timestamp)
    os.makedirs(final_output_dir, exist_ok=True)
    print(f"✅ Output will be saved in: {final_output_dir}")

    # ---------------------------------------
    # ✅ Load ROOT info
    # ---------------------------------------
    root_file = os.path.join(input_dir, "list-roots.json")
    with open(root_file) as f:
        roots_data = json.load(f)
    root_name = roots_data["Roots"][0]["Name"]
    print(f"✅ Root Name: {root_name}")

    # ---------------------------------------
    # ✅ Load Organizational Units
    # ---------------------------------------
    ous_file = os.path.join(input_dir, "list-organizational-units-for-parent.json")
    with open(ous_file) as f:
        ous_data = json.load(f)
    ous_list = ous_data.get("OrganizationalUnits", [])
    print(f"✅ Found {len(ous_list)} Organizational Units.")

    # ---------------------------------------
    # ✅ Load list-accounts.json (ALL accounts)
    # ---------------------------------------
    all_accounts_file = os.path.join(input_dir, "list-accounts.json")
    with open(all_accounts_file) as f:
        all_accounts_data = json.load(f)
    all_accounts_list = all_accounts_data.get("Accounts", [])
    print(f"✅ Loaded {len(all_accounts_list)} total accounts from list-accounts.json.")

    # ---------------------------------------
    # ✅ Load Accounts per OU
    # ---------------------------------------
    accounts_by_ou = {}
    account_files_pattern = os.path.join(input_dir, "list-accounts-for-parent-*.json")
    account_files = glob.glob(account_files_pattern)
    print(f"✅ Found {len(account_files)} account list files for OUs.")

    for account_file in account_files:
        with open(account_file) as f:
            accounts_data = json.load(f)
        filename = os.path.basename(account_file)
        ou_name_raw = filename.replace("list-accounts-for-parent-", "").replace(".json", "")
        ou_name_normalized = normalize(ou_name_raw)
        accounts_list = accounts_data.get("Accounts", [])
        accounts_by_ou[ou_name_normalized] = accounts_list
        print(f"✅ Loaded {len(accounts_list)} accounts for OU Name: {ou_name_raw}")

    # ---------------------------------------
    # ✅ Prepare Data for Tables
    # ---------------------------------------
    # Build mapping of Account ID -> OU Name
    account_id_to_ou = {}
    for ou in ous_list:
        ou_name = ou["Name"]
        clean_ou_name = normalize(ou_name)
        for account in accounts_by_ou.get(clean_ou_name, []):
            account_id_to_ou[account["Id"]] = ou_name

    # ✅ Master Account Table (all accounts)
    master_table_data = []
    for acct in all_accounts_list:
        acct_id = acct.get("Id", "Unknown")
        acct_name = acct.get("Name", "Unknown")
        acct_status = acct.get("Status", "Unknown")
        ou_name = account_id_to_ou.get(acct_id, "None")
        master_table_data.append([acct_name, acct_id, acct_status, ou_name])

    # ✅ OU Table (only assigned)
    ou_table_data = []
    for ou in ous_list:
        ou_name = ou["Name"]
        clean_ou_name = normalize(ou_name)
        for account in accounts_by_ou.get(clean_ou_name, []):
            acct_name = account["Name"]
            acct_id = account.get("Id", "Unknown")
            acct_status = account["Status"]
            ou_table_data.append([ou_name, acct_name, acct_id, acct_status])

    # ---------------------------------------
    # ✅ Write Both Tables
    # ---------------------------------------
    export_table_csv_docx(
        master_table_data,
        ["Account Name", "Account ID", "Status", "OU Name"],
        os.path.join(final_output_dir, "aws_org_all_accounts.csv"),
        os.path.join(final_output_dir, "aws_org_all_accounts.docx"),
        "AWS Organizations - All Accounts (Master List)"
    )

    export_table_csv_docx(
        ou_table_data,
        ["OU Name", "Account Name", "Account ID", "Status"],
        os.path.join(final_output_dir, "aws_org_accounts_by_ou.csv"),
        os.path.join(final_output_dir, "aws_org_accounts_by_ou.docx"),
        "AWS Organizations - Accounts by OU"
    )

    # ---------------------------------------
    # ✅ Build Mermaid Diagram
    # ---------------------------------------
    mermaid_lines = []
    mermaid_lines.append("graph TD")
    mermaid_lines.append(f'  Root["{root_name}"]')

    class_assignments = []

    # Add links from Root to each OU
    for ou in ous_list:
        ou_name = ou["Name"]
        clean_ou_name = normalize(ou_name)
        mermaid_lines.append(f'  Root --> {clean_ou_name}')

    for ou in ous_list:
        ou_name = ou["Name"]
        clean_ou_name = normalize(ou_name)
        mermaid_lines.append(f'  subgraph {clean_ou_name} ["{ou_name}"]')
        accounts = accounts_by_ou.get(clean_ou_name, [])
        for account in accounts:
            acc_name = account["Name"]
            acc_status = account["Status"]
            node_id = normalize(acc_name)
            if normalize(acc_name) == clean_ou_name:
                acc_name_display = f"{acc_name} (Account)"
                node_id = node_id + "Account"
            else:
                acc_name_display = acc_name
            mermaid_lines.append(f'    {node_id}["{acc_name_display} ({acc_status})"]')
            class_assignments.append((node_id, acc_status))
        mermaid_lines.append("  end")

    mermaid_lines.append("")
    mermaid_lines.append("classDef active fill:#28a745,stroke:#333,stroke-width:1px;")
    mermaid_lines.append("classDef suspended fill:#d73a49,stroke:#333,stroke-width:1px;")
    mermaid_lines.append("")
    for node_id, status in class_assignments:
        if status == "ACTIVE":
            mermaid_lines.append(f'class {node_id} active')
        elif status == "SUSPENDED":
            mermaid_lines.append(f'class {node_id} suspended')

    # ✅ Write Mermaid .mmd file
    output_mmd_file = "aws_org_diagram.mmd"
    output_image_file = f"aws_org_diagram.{image_format}"
    output_mmd_path = os.path.join(final_output_dir, output_mmd_file)
    with open(output_mmd_path, "w") as f:
        f.write("\n".join(mermaid_lines))
    print(f"\n✅ Mermaid diagram saved to: {output_mmd_path}")

    # ✅ Render image
    mmdc_path = shutil.which("mmdc")
    if not mmdc_path:
        raise RuntimeError(
            "❗ Mermaid CLI (mmdc) not found in PATH. "
            "Install it with: npm install -g @mermaid-js/mermaid-cli"
        )
    output_image_path = os.path.join(final_output_dir, output_image_file)
    subprocess.run([
        mmdc_path,
        "-i", output_mmd_path,
        "-o", output_image_path,
        "-s", scale
    ], check=True)
    print(f"✅ Diagram image generated at: {output_image_path}")

    return final_output_dir
